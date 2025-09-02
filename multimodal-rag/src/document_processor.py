# src/document_processor.py
import os
import fitz  # PyMuPDF
import requests
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from urllib.parse import urlparse
import mimetypes
# LangChain document loaders
from langchain.document_loaders import (
    PyPDFLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredPowerPointLoader,
    TextLoader,
    CSVLoader,
    UnstructuredHTMLLoader,
    UnstructuredMarkdownLoader,
    UnstructuredExcelLoader,
    WebBaseLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Image processing
from PIL import Image, ExifTags
import pytesseract  # OCR
from bs4 import BeautifulSoup
import zipfile
import json

@dataclass
class ProcessedDocument:
    """처리된 문서 정보"""
    text_docs: List[Document]
    image_docs: List[Document]
    metadata: Dict[str, Any]
    success: bool
    error: Optional[str] = None

class UniversalDocumentProcessor:
    """모든 파일 형식을 지원하는 범용 문서 처리기"""
    
    SUPPORTED_FORMATS = {
        # 문서 형식
        '.pdf': 'PDF 문서',
        '.docx': 'Word 문서',
        '.doc': 'Word 문서 (구버전)',
        '.pptx': 'PowerPoint 프레젠테이션',
        '.ppt': 'PowerPoint 프레젠테이션 (구버전)',
        '.xlsx': 'Excel 스프레드시트',
        '.xls': 'Excel 스프레드시트 (구버전)',
        '.csv': 'CSV 데이터',
        '.txt': '텍스트 파일',
        '.md': 'Markdown 문서',
        '.html': 'HTML 문서',
        '.htm': 'HTML 문서',
        
        # 이미지 형식
        '.jpg': '이미지',
        '.jpeg': '이미지',
        '.png': '이미지',
        '.gif': '이미지',
        '.bmp': '이미지',
        '.tiff': '이미지',
        '.webp': '이미지',
        '.heic': '이미지',
        
        # 압축 파일
        '.zip': '압축 파일',
        
        # 웹 URL
        'http': '웹 문서',
        'https': '웹 문서'
    }
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self.temp_dir = "temp_processing"
        self.ensure_temp_dir()
    
    def ensure_temp_dir(self):
        """임시 디렉토리 생성"""
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)
    
    def process_document(self, source: Union[str, List[str]]) -> ProcessedDocument:
        """문서 또는 문서들을 처리"""
        try:
            if isinstance(source, str):
                return self._process_single_document(source)
            else:
                return self._process_multiple_documents(source)
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=str(e)
            )
    
    def _process_single_document(self, source: str) -> ProcessedDocument:
        """단일 문서 처리"""
        # URL인지 파일 경로인지 판단
        if source.startswith(('http://', 'https://')):
            return self._process_web_document(source)
        else:
            return self._process_file_document(source)
    
    def _process_multiple_documents(self, sources: List[str]) -> ProcessedDocument:
        """여러 문서 처리"""
        all_text_docs = []
        all_image_docs = []
        combined_metadata = {
            "sources": sources,
            "processed_files": [],
            "failed_files": []
        }
        
        for source in sources:
            result = self._process_single_document(source)
            if result.success:
                all_text_docs.extend(result.text_docs)
                all_image_docs.extend(result.image_docs)
                combined_metadata["processed_files"].append(source)
            else:
                combined_metadata["failed_files"].append({
                    "source": source,
                    "error": result.error
                })
        
        return ProcessedDocument(
            text_docs=all_text_docs,
            image_docs=all_image_docs,
            metadata=combined_metadata,
            success=len(all_text_docs) > 0 or len(all_image_docs) > 0
        )
    
    def _process_file_document(self, file_path: str) -> ProcessedDocument:
        """파일 문서 처리"""
        if not os.path.exists(file_path):
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"파일을 찾을 수 없습니다: {file_path}"
            )
        
        file_ext = Path(file_path).suffix.lower()
        
        # 파일 형식별 처리
        if file_ext == '.pdf':
            return self._process_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._process_word_document(file_path)
        elif file_ext in ['.pptx', '.ppt']:
            return self._process_powerpoint(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self._process_excel(file_path)
        elif file_ext == '.csv':
            return self._process_csv(file_path)
        elif file_ext == '.txt':
            return self._process_text(file_path)
        elif file_ext == '.md':
            return self._process_markdown(file_path)
        elif file_ext in ['.html', '.htm']:
            return self._process_html_file(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.heic']:
            return self._process_image_file(file_path)
        elif file_ext == '.zip':
            return self._process_zip_file(file_path)
        else:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"지원하지 않는 파일 형식: {file_ext}"
            )
    
    def _process_pdf(self, pdf_path: str) -> ProcessedDocument:
        """PDF 처리 (기존 로직 개선)"""
        text_docs = []
        image_docs = []
        
        # 1. 텍스트 추출
        try:
            loader = PyPDFLoader(pdf_path)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
        except Exception as e:
            pass  # 텍스트 추출 실패해도 이미지는 처리
        
        # 2. 이미지 추출
        try:
            images_info = self._extract_pdf_images(pdf_path)
            for img_info in images_info:
                # 이미지 설명은 외부에서 추가 (image_analyzer.py에서)
                image_doc = Document(
                    page_content="",  # 빈 상태로, 나중에 분석 결과 추가
                    metadata={
                        "source": pdf_path,
                        "page": img_info['page'],
                        "type": "image",
                        "image_path": img_info['path'],
                        "filename": img_info['filename'],
                        "original_format": "pdf"
                    }
                )
                image_docs.append(image_doc)
        except Exception as e:
            pass
        
        return ProcessedDocument(
            text_docs=text_docs,
            image_docs=image_docs,
            metadata={
                "source": pdf_path,
                "type": "pdf",
                "text_chunks": len(text_docs),
                "images": len(image_docs)
            },
            success=len(text_docs) > 0 or len(image_docs) > 0
        )
    
    def _process_word_document(self, file_path: str) -> ProcessedDocument:
        """Word 문서 처리"""
        text_docs = []
        image_docs = []
        
        try:
            # 텍스트 추출
            loader = UnstructuredWordDocumentLoader(file_path)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            # Word 문서에서 이미지 추출
            images_info = self._extract_word_images(file_path)
            for img_info in images_info:
                image_doc = Document(
                    page_content="",
                    metadata={
                        "source": file_path,
                        "type": "image",
                        "image_path": img_info['path'],
                        "filename": img_info['filename'],
                        "original_format": "docx"
                    }
                )
                image_docs.append(image_doc)
                
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"Word 문서 처리 오류: {str(e)}"
            )
        
        return ProcessedDocument(
            text_docs=text_docs,
            image_docs=image_docs,
            metadata={
                "source": file_path,
                "type": "word",
                "text_chunks": len(text_docs),
                "images": len(image_docs)
            },
            success=True
        )
    
    def _process_powerpoint(self, file_path: str) -> ProcessedDocument:
        """PowerPoint 프레젠테이션 처리"""
        try:
            loader = UnstructuredPowerPointLoader(file_path)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            # PPT 이미지 추출 (복잡한 구현 필요)
            image_docs = []  # 일단 텍스트만
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=image_docs,
                metadata={
                    "source": file_path,
                    "type": "powerpoint",
                    "text_chunks": len(text_docs),
                    "images": len(image_docs)
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"PowerPoint 처리 오류: {str(e)}"
            )
    
    def _process_excel(self, file_path: str) -> ProcessedDocument:
        """Excel 파일 처리"""
        try:
            loader = UnstructuredExcelLoader(file_path)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=[],  # Excel은 보통 차트가 있지만 복잡한 추출 필요
                metadata={
                    "source": file_path,
                    "type": "excel",
                    "text_chunks": len(text_docs)
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"Excel 처리 오류: {str(e)}"
            )
    
    def _process_csv(self, file_path: str) -> ProcessedDocument:
        """CSV 파일 처리"""
        try:
            loader = CSVLoader(file_path)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=[],
                metadata={
                    "source": file_path,
                    "type": "csv",
                    "text_chunks": len(text_docs)
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"CSV 처리 오류: {str(e)}"
            )
    
    def _process_text(self, file_path: str) -> ProcessedDocument:
        """텍스트 파일 처리"""
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=[],
                metadata={
                    "source": file_path,
                    "type": "text",
                    "text_chunks": len(text_docs)
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"텍스트 파일 처리 오류: {str(e)}"
            )
    
    def _process_markdown(self, file_path: str) -> ProcessedDocument:
        """Markdown 파일 처리"""
        try:
            loader = UnstructuredMarkdownLoader(file_path)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=[],
                metadata={
                    "source": file_path,
                    "type": "markdown",
                    "text_chunks": len(text_docs)
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"Markdown 처리 오류: {str(e)}"
            )
    
    def _process_html_file(self, file_path: str) -> ProcessedDocument:
        """HTML 파일 처리"""
        try:
            loader = UnstructuredHTMLLoader(file_path)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            # HTML에서 이미지 링크 추출
            image_docs = self._extract_html_images(file_path)
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=image_docs,
                metadata={
                    "source": file_path,
                    "type": "html",
                    "text_chunks": len(text_docs),
                    "images": len(image_docs)
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"HTML 처리 오류: {str(e)}"
            )
    
    def _process_image_file(self, image_path: str) -> ProcessedDocument:
        """단일 이미지 파일 처리"""
        try:
            # OCR로 텍스트 추출
            text_docs = []
            try:
                image = Image.open(image_path)
                ocr_text = pytesseract.image_to_string(image, lang='kor+eng')
                if ocr_text.strip():
                    text_doc = Document(
                        page_content=f"[OCR 추출 텍스트] {ocr_text}",
                        metadata={
                            "source": image_path,
                            "type": "ocr_text",
                            "original_format": "image"
                        }
                    )
                    text_docs.append(text_doc)
            except:
                pass  # OCR 실패해도 이미지는 처리
            
            # 이미지 문서 생성
            image_doc = Document(
                page_content="",  # 나중에 vision model로 분석
                metadata={
                    "source": image_path,
                    "type": "image",
                    "image_path": image_path,
                    "filename": os.path.basename(image_path),
                    "original_format": "image"
                }
            )
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=[image_doc],
                metadata={
                    "source": image_path,
                    "type": "image",
                    "text_chunks": len(text_docs),
                    "images": 1
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"이미지 처리 오류: {str(e)}"
            )
    
    def _process_web_document(self, url: str) -> ProcessedDocument:
        """웹 문서 처리"""
        try:
            loader = WebBaseLoader(url)
            documents = loader.load()
            text_docs = self.text_splitter.split_documents(documents)
            
            # 웹페이지에서 이미지 추출
            image_docs = self._extract_web_images(url)
            
            return ProcessedDocument(
                text_docs=text_docs,
                image_docs=image_docs,
                metadata={
                    "source": url,
                    "type": "web",
                    "text_chunks": len(text_docs),
                    "images": len(image_docs)
                },
                success=True
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"웹 문서 처리 오류: {str(e)}"
            )
    
    def _process_zip_file(self, zip_path: str) -> ProcessedDocument:
        """압축 파일 처리"""
        try:
            all_text_docs = []
            all_image_docs = []
            
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                extract_path = os.path.join(self.temp_dir, "extracted")
                zip_ref.extractall(extract_path)
                
                # 압축 해제된 파일들을 재귀적으로 처리
                for root, dirs, files in os.walk(extract_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        result = self._process_file_document(file_path)
                        if result.success:
                            all_text_docs.extend(result.text_docs)
                            all_image_docs.extend(result.image_docs)
            
            return ProcessedDocument(
                text_docs=all_text_docs,
                image_docs=all_image_docs,
                metadata={
                    "source": zip_path,
                    "type": "zip",
                    "text_chunks": len(all_text_docs),
                    "images": len(all_image_docs)
                },
                success=len(all_text_docs) > 0 or len(all_image_docs) > 0
            )
        except Exception as e:
            return ProcessedDocument(
                text_docs=[],
                image_docs=[],
                metadata={},
                success=False,
                error=f"압축 파일 처리 오류: {str(e)}"
            )
    
    def _extract_pdf_images(self, pdf_path: str) -> List[Dict[str, Any]]:
        """PDF 이미지 추출 (기존 로직)"""
        doc = fitz.open(pdf_path)
        image_info = []
        
        try:
            for page_num in range(doc.page_count):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:
                            img_filename = f"pdf_page_{page_num+1}_img_{img_index+1}.png"
                            img_path = os.path.join(self.temp_dir, img_filename)
                            
                            pix.save(img_path)
                            
                            image_info.append({
                                "path": img_path,
                                "page": page_num + 1,
                                "index": img_index + 1,
                                "filename": img_filename
                            })
                        
                        pix = None
                    except Exception as e:
                        continue
        finally:
            doc.close()
        
        return image_info
    
    def _extract_word_images(self, docx_path: str) -> List[Dict[str, Any]]:
        """Word 문서에서 이미지 추출"""
        image_info = []
        try:
            # python-docx로 이미지 추출
            from docx import Document as DocxDocument
            
            doc = DocxDocument(docx_path)
            
            for i, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    img_data = rel.target_part.blob
                    img_filename = f"word_img_{i+1}.png"
                    img_path = os.path.join(self.temp_dir, img_filename)
                    
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                    
                    image_info.append({
                        "path": img_path,
                        "filename": img_filename
                    })
        except Exception as e:
            pass
        
        return image_info
    
    def _extract_html_images(self, html_path: str) -> List[Document]:
        """HTML 파일에서 이미지 링크 추출"""
        image_docs = []
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            for img in soup.find_all('img'):
                img_src = img.get('src')
                if img_src:
                    image_doc = Document(
                        page_content="",
                        metadata={
                            "source": html_path,
                            "type": "image",
                            "image_url": img_src,
                            "alt_text": img.get('alt', ''),
                            "original_format": "html"
                        }
                    )
                    image_docs.append(image_doc)
        except Exception as e:
            pass
        
        return image_docs
    
    def _extract_web_images(self, url: str) -> List[Document]:
        """웹페이지에서 이미지 추출"""
        image_docs = []
        try:
            response = requests.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            for img in soup.find_all('img')[:10]:  # 최대 10개
                img_src = img.get('src')
                if img_src:
                    # 상대 경로를 절대 경로로 변환
                    if img_src.startswith('/'):
                        base_url = f"{urlparse(url).scheme}://{urlparse(url).netloc}"
                        img_src = base_url + img_src
                    elif not img_src.startswith('http'):
                        img_src = url + '/' + img_src
                    
                    image_doc = Document(
                        page_content="",
                        metadata={
                            "source": url,
                            "type": "image",
                            "image_url": img_src,
                            "alt_text": img.get('alt', ''),
                            "original_format": "web"
                        }
                    )
                    image_docs.append(image_doc)
        except Exception as e:
            pass
        
        return image_docs
    
    def get_supported_formats(self) -> Dict[str, str]:
        """지원하는 파일 형식 목록 반환"""
        return self.SUPPORTED_FORMATS.copy()
    
    def is_supported(self, file_path_or_url: str) -> bool:
        """파일이 지원되는 형식인지 확인"""
        if file_path_or_url.startswith(('http://', 'https://')):
            return True
        
        file_ext = Path(file_path_or_url).suffix.lower()
        return file_ext in self.SUPPORTED_FORMATS